import os
import json
import requests
import html
import re
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from config.settings import OPENCODE_API_KEY, OPENCODE_API_URL, RESUMES_DIR

def clean_and_escape(text):
    """
    Escapes special characters (&, <, >) to avoid ReportLab Paragraph parsing exceptions
    and safely converts Markdown bold/italic accents to Paragraph inline styling.
    """
    if not text:
        return ""
    # XML Escape raw characters safely
    text = html.escape(text)
    # Restore the escaped &amp; to just &amp; (since html.escape might double escape if we pass already-escaped characters)
    text = text.replace("&amp;amp;", "&amp;")
    
    # Map double asterisks markdown to HTML Bold tags
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'__(.*?)__', r'<b>\1</b>', text)
    
    # Map single asterisks markdown to HTML Italic tags
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    text = re.sub(r'_(.*?)_', r'<i>\1</i>', text)
    
    # Standard XML entities restore
    text = text.replace("&amp;bull;", "&bull;").replace("&amp;middot;", "&middot;")
    return text

def call_opencode_api(user_data, model_id):
    """
    Calls the OpenCode API utilizing the active selected model
    to generate an elegant, ATS-friendly professional resume based on user provided data.
    """
    if not OPENCODE_API_KEY:
        # Fallback raw builder if no API key is set
        return generate_static_resume_text(user_data)

    system_prompt = (
        "You are an elite, professional CV writer and HR recruiter specializing in ATS-friendly formatting, "
        "professional typography, and high-impact industry verbiage. "
        "Generate a complete, fully expanded professional Resume/CV matching the user spec. "
        "Choose an extremely professional, clean tone. Avoid AI phrases or generic buzzwords. "
        "Return the CV under these EXACT Markdown sections so the system can parse it if needed: "
        "\n# [FULL_NAME]"
        "\n## Contact Info: [PHONE_NUMBER] | [EMAIL] | [ADDRESS]"
        "\n## Summary or Objective\n[SUMMARY_TEXT]"
        "\n## Professional Skills\n[SKILLS_LIST]"
        "\n## Professional Experience\n[EXPERIENCE_LIST]"
        "\n## Education Profile\n[EDUCATION_LIST]"
        "\n## Core Projects\n[PROJECTS_LIST]"
        "\n## Certifications\n[CERTIFICATIONS_LIST]"
        "\n## Languages\n[LANGUAGES_LIST]"
        "\n## Hobbies & Interests\n[HOBBIES_LIST]"
        "Formatting instructions: Expand everything to look impressive. Format lists beautifully with bullet points."
    )

    prompt = (
        f"Generate a customized comprehensive executive resume based on this collected profile:\n"
        f"Name: {user_data.get('name')}\n"
        f"Phone: {user_data.get('phone')}\n"
        f"Email: {user_data.get('email')}\n"
        f"Address: {user_data.get('address')}\n"
        f"Career Objective: {user_data.get('objective')}\n"
        f"Skills: {user_data.get('skills')}\n"
        f"Education Profile: {user_data.get('education')}\n"
        f"Professional Experience: {user_data.get('experience')}\n"
        f"Key Projects: {user_data.get('projects')}\n"
        f"Certifications: {user_data.get('certifications')}\n"
        f"Languages spoken: {user_data.get('languages')}\n"
        f"Hobbies & Interests: {user_data.get('hobbies')}\n"
    )

    headers = {
        "Authorization": f"Bearer {OPENCODE_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model_id,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }

    try:
        response = requests.post(OPENCODE_API_URL, json=payload, headers=headers, timeout=60)
        if response.status_code == 200:
            result = response.json()
            return result["choices"][0]["message"]["content"]
        else:
            # Fallback
            return generate_static_resume_text(user_data, f"API Error: Status {response.status_code}")
    except Exception as e:
        return generate_static_resume_text(user_data, f"Connection Failure: {str(e)}")

def generate_static_resume_text(u, fallback_msg=""):
    """
    Simple fallback text compilation if the AI provider cannot be resolved.
    """
    text = f"# {u.get('name', 'QuickCV User')}\n"
    text += f"Contact: {u.get('phone')} | {u.get('email')} | {u.get('address')}\n\n"
    if fallback_msg:
        text += f"*Generated with System Fallback* — {fallback_msg}\n\n"
    text += f"## Career Objective\n{u.get('objective', 'To excel in a challenging environment.')}\n\n"
    text += f"## Technical & Professional Skills\n{u.get('skills', '').replace(',', ', ')}\n\n"
    text += f"## Work Experience\n{u.get('experience', '')}\n\n"
    text += f"## Educational Background\n{u.get('education', '')}\n\n"
    text += f"## Strategic Projects\n{u.get('projects', '')}\n\n"
    text += f"## Certifications\n{u.get('certifications', '')}\n\n"
    text += f"## Languages\n{u.get('languages', '')}\n\n"
    text += f"## Hobbies\n{u.get('hobbies', '')}\n"
    return text

def parse_markdown_to_sections(text):
    """
    Parses resume markdown into dictionary sections for programmatic styling
    """
    lines = text.split("\n")
    sections = {
        "name": "",
        "contact": "",
        "sections": []
    }
    
    current_sec_title = None
    current_sec_body = []
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            sections["name"] = stripped[2:].strip()
        elif stripped.startswith("## Contact Info:") or "Contact Info:" in stripped:
            content = stripped.replace("## Contact Info:", "").replace("Contact Info:", "").strip()
            sections["contact"] = content
        elif stripped.startswith("## ") or stripped.startswith("### "):
            if current_sec_title:
                sections["sections"].append({
                    "title": current_sec_title,
                    "body": "\n".join(current_sec_body).strip()
                })
            # Start new section
            current_sec_title = stripped.lstrip("#").strip()
            current_sec_body = []
        else:
            if current_sec_title:
                current_sec_body.append(line)
            else:
                # Top header fallback
                if not sections["name"] and not stripped.startswith("#"):
                    sections["name"] = stripped
                else:
                    if sections["contact"] == "":
                        sections["contact"] = stripped
                    else:
                        sections["contact"] += f" | {stripped}"
                        
    if current_sec_title:
        sections["sections"].append({
            "title": current_sec_title,
            "body": "\n".join(current_sec_body).strip()
        })
        
    if not sections["name"]:
        sections["name"] = "Professional Resume"
        
    return sections

def create_resume_files(user_id, resume_text, unique_id):
    """
    Compiles PDF and DOCX files representing the given resume markdown.
    Stores files in /generated_resumes/ using unique identifiers.
    Returns (pdf_path, docx_path) filenames.
    """
    parsed = parse_markdown_to_sections(resume_text)
    
    pdf_filename = f"resume_{user_id}_{unique_id}.pdf"
    docx_filename = f"resume_{user_id}_{unique_id}.docx"
    
    pdf_path = os.path.join(RESUMES_DIR, pdf_filename)
    docx_path = os.path.join(RESUMES_DIR, docx_filename)
    
    # 1. GENERATE PDF WITH REPORTLAB
    try:
        # Create standard Letter page template
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Color Palette Design: Slate Black & Subtle Navy Accent
        primary_color = colors.HexColor("#1A202C")  # Slate Gray Black
        accent_color = colors.HexColor("#2B6CB0")   # Indigo Blue accent
        text_color = colors.HexColor("#2D3748")     # Soft Charcoal
        
        # Define Custom Styles
        title_style = ParagraphStyle(
            'CVTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=28,
            textColor=primary_color,
            alignment=1, # Centered
            spaceAfter=6
        )
        
        contact_style = ParagraphStyle(
            'CVContact',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9.5,
            leading=13,
            textColor=accent_color,
            alignment=1, # Centered
            spaceAfter=15
        )
        
        heading_style = ParagraphStyle(
            'CVHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=13,
            leading=16,
            textColor=primary_color,
            spaceBefore=14,
            spaceAfter=4,
            keepWithNext=True
        )
        
        body_style = ParagraphStyle(
            'CVBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=text_color,
            spaceAfter=6
        )
        
        # Header components
        story.append(Paragraph(clean_and_escape(parsed["name"]), title_style))
        story.append(Paragraph(clean_and_escape(parsed["contact"]), contact_style))
        
        # Horizontal line divider
        divider_table = Table([[""]], colWidths=[504], rowHeights=[1.5])
        divider_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), accent_color),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0),
            ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        story.append(divider_table)
        story.append(Spacer(1, 10))
        
        # Render each parsed section
        for sec in parsed["sections"]:
            # Heading Section
            story.append(Paragraph(clean_and_escape(sec["title"]), heading_style))
            
            # Left accent borderline
            sec_divider = Table([[""]], colWidths=[504], rowHeights=[1])
            sec_divider.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#E2E8F0")),
                ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
            ]))
            story.append(sec_divider)
            story.append(Spacer(1, 4))
            
            # Content Paragraphs
            body_lines = sec["body"].split("\n")
            for bline in body_lines:
                if not bline.strip():
                    continue
                # Style bullet points nicely
                if bline.strip().startswith("-") or bline.strip().startswith("*"):
                    cleared = bline.strip().lstrip("-*").strip()
                    story.append(Paragraph(f"&bull; {clean_and_escape(cleared)}", body_style))
                else:
                    story.append(Paragraph(clean_and_escape(bline), body_style))
            story.append(Spacer(1, 8))
            
        doc.build(story)
    except Exception as e:
        # Fallback raw text PDF builder in case of parser crashes
        try:
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph("QuickCV Document Generation Fallback", styles['Heading1'])]
            for line in resume_text.split("\n"):
                story.append(Paragraph(clean_and_escape(line), styles['Normal']))
            doc.build(story)
        except:
            pass

    # 2. GENERATE DOCX WITH python-docx
    try:
        doc = Document()
        
        # Set margins to 1 inch
        for section in doc.sections:
            section.top_margin = Inches_fallback(1)
            section.bottom_margin = Inches_fallback(1)
            section.left_margin = Inches_fallback(1)
            section.right_margin = Inches_fallback(1)
            
        # CV Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(parsed["name"])
        title_run.font.name = 'Arial'
        title_run.font.size = Pt_fallback(24)
        title_run.bold = True
        title_para.alignment = 1 # Centered
        
        # Contact details
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(parsed["contact"])
        contact_run.font.name = 'Arial'
        contact_run.font.size = Pt_fallback(10)
        contact_para.alignment = 1 # Centered
        
        # Iterate over sections
        for sec in parsed["sections"]:
            # Header Section
            h_para = doc.add_paragraph()
            h_para.paragraph_format.space_before = Pt_fallback(12)
            h_para.paragraph_format.space_after = Pt_fallback(3)
            h_run = h_para.add_run(sec["title"].upper())
            h_run.font.name = 'Arial'
            h_run.font.size = Pt_fallback(12)
            h_run.bold = True
            
            # Content Paragraphs
            body_lines = sec["body"].split("\n")
            for bline in body_lines:
                if not bline.strip():
                    continue
                if bline.strip().startswith("-") or bline.strip().startswith("*"):
                    cleared = bline.strip().lstrip("-*").strip()
                    try:
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_after = Pt_fallback(3)
                        p_run = p.add_run(cleared)
                        p_run.font.name = 'Arial'
                        p_run.font.size = Pt_fallback(10.5)
                    except:
                        # Fallback if 'List Bullet' template style is missing in Word configuration
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt_fallback(3)
                        p_run = p.add_run(f"• {cleared}")
                        p_run.font.name = 'Arial'
                        p_run.font.size = Pt_fallback(10.5)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt_fallback(4)
                    p_run = p.add_run(bline)
                    p_run.font.name = 'Arial'
                    p_run.font.size = Pt_fallback(10.5)
                    
        doc.save(docx_path)
    except Exception as e:
        # Fallback very simple DOCX write
        try:
            doc = Document()
            doc.add_heading("CV Document", 0)
            doc.add_paragraph(resume_text)
            doc.save(docx_path)
        except:
            pass

    return pdf_filename, docx_filename

def Inches_fallback(inches):
    """Fallback manual Inches setup for python-docx"""
    from docx.shared import Inches
    return Inches(inches)

def Pt_fallback(pt):
    """Fallback manual Pt setup for python-docx"""
    from docx.shared import Pt
    return Pt(pt)
